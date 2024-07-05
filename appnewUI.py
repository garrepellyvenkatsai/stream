import streamlit as st
import pandas as pd
from io import StringIO
from langchain.document_loaders import DirectoryLoader
from langchain.text_splitter import CharacterTextSplitter
from PyPDF2 import PdfReader
from langchain.docstore.document import Document
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from dotenv import load_dotenv
import os
import time
import base64
from xml.dom.minidom import parseString
from docx import Document
from tempfile import NamedTemporaryFile
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from streamlit_chat import message
from langchain.prompts import PromptTemplate
from PIL import Image
from streamlit_extras.stylable_container import stylable_container
from langchain.chains import ConversationalRetrievalChain
from langchain.document_loaders import PyPDFLoader
from streamlit_float import *
from pdf2image import convert_from_bytes
from langchain.chains import ConversationalRetrievalChain
from msal import ConfidentialClientApplication
from langchain.chat_models import AzureChatOpenAI
import openai
import time

#customise frontend
st.set_page_config(page_title="Chat With Doc")
logo = Image.open('logo.png')
# Create two columns
col1, col2 = st.columns([1, 4])

# Display the logo in the first column
with col1:
    st.image(logo, width=100)

# Display the title in the second column
with col2:
    st.title("Chat With Doc")



load_dotenv()

st.markdown(
    """
    <style>
    .title-container {
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 10px;
        margin-bottom: 20px;
    }
    .title-container img {
        width: 50px;
    }
    .title-container h1 {
        margin: 0;
    }
    .main-container {
        display: flex;
        flex-direction: column;
        height: 100vh;
    }
    
    .input-container {
        max-width: 800px;
        margin: auto;
        padding: 10px;
        position: fixed;
        bottom: 0;
        width: 100%;
        background-color: #DCF8C6;  
        border-top: 1px solid #ddd;
        z-index: 1000;
    }
    .chat-message {
        padding: 10px;
        border-radius: 10px;
        margin: 5px 0;
        width: fit-content;
        max-width: 80%;
    }
    .user-message {
        background-color: #DCF8C6;
        align-self: flex-end;
        text-align: right;
        margin-left: auto;
    }
    .assistant-message {
        background-color: #E8E8E8;
        align-self: flex-start;
        text-align: left;
    }
    </style>
    """,
    unsafe_allow_html=True
)
os.environ["OPENAI_API_TYPE"]= os.getenv("OPENAI_API_TYPE")
os.environ["OPENAI_API_VERSION"]= os.getenv("OPENAI_API_VERSION")
os.environ['OPENAI_AZURE_CLIENT_SECRET']= os.getenv("OPENAI_AZURE_CLIENT_SECRET")
os.environ['OPENAI_AZURE_CLIENT_ID']= os.getenv("OPENAI_AZURE_CLIENT_ID")

# below generated2 token is valid for 1 hr from its creation.
# to generate the key/ secret
@st.cache_data(ttl=2700)
def getSecret():
    # orgg_constants\n
    print('Running get secret')
    tenant_id = "7c917db0-71f2-438e-9554-388ffcab8764"
    auth = f"https://login.microsoftonline.com/{tenant_id}"
    a_cid = "04b07795-8ddb-461a-bbee-02f9e1bf7b46"

    # Providing Azure Client/App ID and Secret to Variables
    client_secret = os.environ['OPENAI_AZURE_CLIENT_SECRET']
    client_id=os.environ['OPENAI_AZURE_CLIENT_ID']
    scope_list = [client_id+'/.default']

    # Getting JWT Token Using MSAL Methods for OpenAI Authentication\n
    app = ConfidentialClientApplication(
        client_id=client_id,
        client_credential=client_secret,
        authority=auth
    )
    result = app.acquire_token_for_client(scopes=scope_list)
    access_token = result.get("access_token")
    # print(access_token) ## JWT Token
    # print('Access token:', access_token)
    print('Access token fetched. Returning.')
    return access_token

# Obtain the access token
secret1= getSecret()

# #phase-1
# embeddings = OpenAIEmbeddings(
#         deployment="snd-014-text-embedding-ada-002",
#         model="text-embedding-ada-002",
#         openai_api_base="https://openai-014-cgs.openai.azure.com/",
#         openai_api_type="azure",
#     )

# #phase-2
app_shortform = "uc24"  #will be provided to use case team in onboarding email
environment = "appdev"    # Requested Environment (appdev/apptest/prod)
# model = "text-embedding-ada-002"           # Model Name in api_base will be provided to usecase team in onboarding email.

embed_deployment_name = "ada002_2" 
base_api = "https://openaiapim-prd-01-weu-002-apim.azure-api.net/paybot/appdev/gpt35/openai/deployments/gpt35_0301/chat/completions?api-version=2023-05-15"
# embed_api_base = base_api + "/"+app_shortform+"/"+environment+"/"+model, ### EX: https://openaiapim-prd-01-weu-002-apim.azure-api.net/uc01/prod/adaxx"
# embed_api_base= "https://openaiapim-prd-01-weu-002-apim.azure-api.net/paybot/appdev/ada"
embed_api_base= 'https://openaiapim-prd-01-weu-002-apim.azure-api.net/paybot2/appdev/ada'

embeddings = OpenAIEmbeddings(
    openai_api_key= secret1,#os.getenv('OPENAI_API_KEY'),
    openai_api_type=os.getenv('OPENAI_API_TYPE'),
    openai_api_base=embed_api_base,
    deployment=embed_deployment_name
    )

# Float feature initialization
float_init()

def add_heading():
    st.markdown(
        """
        <style>
            [data-testid="stSidebarNav"]::before {
                content: "Payments CoPilot";
                margin-left: 15px;
                margin-top: 15px;
                font-size: 40px;
                position: relative;
                top: 50px;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

# display Payments CoPilot in the sideBar


def show_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

def show_docx(file_path):
    doc = Document(file_path)
    content = ""
    for para in doc.paragraphs:
        content += para.text + "\n"
    st.text_area("Document Content", content, height=800)

def show_txt(file_path):
    with open(file_path, "r") as f:
        content = f.read()
    st.text_area("Document Content", content, height=800)

def show_xml(file_path):
    content = file_path.decode("utf-8")
    pretty_xml = parseString(content).toprettyxml()
    st.text_area("XML Content", pretty_xml, height=800)

def show_html(file_path):
    content = file_path.decode("utf-8")
    st.markdown(content, unsafe_allow_html=True)

def show_bpmn(file_path):
    content = file_path.decode("utf-8")
    pretty_bpmn = parseString(content).toprettyxml()
    st.text_area("BPMN Content", pretty_bpmn, height=800)

import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
import streamlit as st
from transformers import BertTokenizer, BertModel
#import torch

class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

class CharacterTextSplitter:
    def __init__(self, chunk_size=1000, chunk_overlap=0):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_documents(self, docs):
        chunks = []
        for doc in docs:
            text = doc.page_content
            for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
                chunks.append(Document(page_content=text[i:i+self.chunk_size], metadata=doc.metadata))
        return chunks

# Function to parse BPMN file and extract elements
def parse_bpmn(file):
    tree = ET.parse(file)
    root = tree.getroot()
    ns = {'bpmn': 'http://www.omg.org/spec/BPMN/20100524/MODEL'}

    elements = []
    for process in root.findall('bpmn:process', ns):
        for elem in process:
            tag = elem.tag.split('}')[1]  # Get tag name without namespace
            if tag in ['task', 'startEvent', 'endEvent', 'exclusiveGateway']:
                elements.append((tag, elem.get('name')))
    return elements

# Function to convert elements to a textual format
def elements_to_text(elements):
    return ' '.join([f"{tag}:{name}" for tag, name in elements if name])

# Function to generate embeddings using BERT
def get_embeddings(text):
    tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
    model = BertModel.from_pretrained('bert-base-uncased')

    inputs = tokenizer(text, return_tensors='pt')
    outputs = model(**inputs)
    embeddings = outputs.last_hidden_state.mean(dim=1)  # Mean pooling
    return embeddings

def doc_preprocessing(uploaded_file):
    if uploaded_file.type == "application/pdf":
        # PDF processing
        reader = PdfReader(uploaded_file)
        docs = []
        i = 1
        for page in reader.pages:
            docs.append(Document(page_content=page.extract_text(), metadata={'page': i}))
            i += 1
        st.success("Loaded PDF!")
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # DOCX processing
        docx_file = DocxDocument(uploaded_file)
        text = ""
        for paragraph in docx_file.paragraphs:
            text += paragraph.text
        docs = [Document(page_content=text, metadata={'page': 1})]
        st.success("Loaded DOCX!")
    elif uploaded_file.type == "text/plain":
        # TXT processing
        with uploaded_file.open() as f:
            text = f.read().decode("utf-8")
        docs = [Document(page_content=text, metadata={'page': 1})]
        st.success("Loaded TXT!")
    elif uploaded_file.type == "text/html":
        # HTML processing
        with uploaded_file.open() as f:
            soup = BeautifulSoup(f, 'html.parser')
        text = soup.get_text()
        docs = [Document(page_content=text, metadata={'page': 1})]
        st.success("Loaded HTML!")
    elif uploaded_file.type == "application/xml" or uploaded_file.type == "text/xml" or uploaded_file.name.endswith(".bpmn"):
        # BPMN processing
        elements = parse_bpmn(uploaded_file)
        text_representation = elements_to_text(elements)
        docs = [Document(page_content=text_representation, metadata={'page': 1})]
        st.success("Loaded BPMN!")
    else:
        st.error("Unsupported file format!")
        return []

    # Initialising chunk size to 1000 and chunk overlap to 0
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    # Splitting the documents into chunks and returning them
    docs_split = text_splitter.split_documents(docs)
    print('Docs length:', len(docs))
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    # Splitting the documents into chunks 
    docs_split = text_splitter.split_documents(docs)
    print('Done splitting the docs!')
    # Check docs_split length-the type of it
    print('Doc split length:', len(docs_split))
    return docs_split

#uploaded_file = st.file_uploader("Upload a file", type=["pdf", "docx", "txt", "html", "xml", "bpmn"])
#####-> processing in parallel, time display to embed chunks, progress bar as no. of steps, be more deterministic
def add_custom_css():
    st.markdown(
        """
        <style>
        /* Style for the progress bar */
        .stProgress > div > div > div > div {
            background-color: #4CAF50; /* Green color */
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    
def generateEmbeddings(texts, persist_dir):
    add_custom_css()
    # to handle rate limit error
    texts_len= len(texts)
    # diplay "please wait" in the UI while embeddings are getting generated2
    with st.spinner('Please wait, Setting it up!'):
        print("Starting the generation of embeddings")
        # st.write("Starting the generation of embeddings")
        progress_placeholder = st.empty()
        time_placeholder = st.empty()
        progress_bar = progress_placeholder.progress(0)

        for i in range(0,texts_len,10):
            
            start_time = time.time()
            print(persist_dir)
            db = Chroma.from_documents(documents=texts[i:i+10], embedding=embeddings, persist_directory=persist_dir)
            print(len(texts[i:i + 10]))
            db.persist()
            db = None
            end_time = time.time()

            elapsed_time = end_time - start_time
            time_placeholder.write(f"Time taken to generate embeddings for chunk {i//10 + 1}: {elapsed_time:.2f} seconds")

            progress_percentage = min((i + 10) / texts_len, 1.0)
            progress_bar.progress(progress_percentage)

            if i + 10 < texts_len:
                time.sleep(4)  # To handle rate limit error
        progress_placeholder.empty()
        time_placeholder.empty()  
        st.write("All chunks processed.")
            # time.sleep(60)
            # print('---')
    # when is it made true before
    # st.session_state.generateEmbs= 'False' 
    # when is it made true after?
    # generateEmbs= False
    

def search_pdf(db, secret1):
    prompt_template = """You are a subject matter expert (SME) assistant for a payments bank in UK and answer the question from the document uploaded first. If information is not enough then take use of global internet search.
    If the context is not relevant, please say "I'm sorry, but I currently don't have enough context to answer your question."

    Context: {context}

    Question: {question}
    """
    PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain_type_kwargs = {"prompt": PROMPT}
    # model = ChatOpenAI(engine="snd-014-gpt-35-turbo", model="gpt-3.5")
    # phase-2 api
    ### Defining the API_Base for Chat Model
    os.environ['OPENAI_GPT_DEPLOYMENT_NAME']= 'gpt4_1106-Preview'
    chatgpt_deployment_name = os.environ['OPENAI_GPT_DEPLOYMENT_NAME'] ##Deployment Name provided in Email for each chat model
    # suffix = 'paybot/appdev/gpt35' # Provided suffix information for each model
    # chat_api_base = 'https://openaiapim-prd-01-weu-002-apim.azure-api.net/' + suffix
    chat_api_base= 'https://openaiapim-prd-01-weu-002-apim.azure-api.net/paybot2/appdev/gpt4'

    # print('Inside search_db b4 chatOpenAI sec:', secret1)
    model= AzureChatOpenAI(
        openai_api_key= secret1,#os.getenv('OPENAI_API_KEY'),
        openai_api_type= "azure_ad",
        openai_api_base=chat_api_base, 
        openai_api_version = "2023-05-15",
        deployment_name=chatgpt_deployment_name,
    )
    qa_chain = RetrievalQA.from_chain_type(model, retriever=db.as_retriever(search_kwargs={'k': 7}), return_source_documents=True, chain_type_kwargs= chain_type_kwargs)
    # qa_chain = ConversationalRetrievalChain.from_llm(model, retriever=db.as_retriever(search_kwargs={'k': 7}), return_source_documents=True)
    return qa_chain
#new UI addition here ->
if "generated2" not in st.session_state:
    st.session_state["generated2"] = ["Ask away any questions you have about the Document."]
if "past2" not in st.session_state:
    st.session_state["past2"] = ["Hey there!"]

# Function to display the conversation
def display_conversation(history):
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    for i in range(len(history["generated2"])):
        st.markdown(f'<div class="chat-message user-message">{history["past2"][i]}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="chat-message assistant-message">{history["generated2"][i]}</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# Function to handle user input
def handle_input(qa):
    with st.container():
        user_input = st.text_input("", placeholder="Enter your prompt here", key="input")
        if st.button("Send"):
            if user_input:
                st.session_state["past2"].append(user_input)
                # Call your QA model to get the response
                response = qa({"question": user_input, "chat_history": st.session_state["past2"]})
                st.session_state["generated2"].append(response["answer"])
                st.experimental_rerun()

def chatInterface(qa):
    # Display the chat interface
    with st.container():
        display_conversation(st.session_state)
        st.markdown('<div class="input-container">', unsafe_allow_html=True)
        handle_input(qa)
        st.markdown('</div>', unsafe_allow_html=True)

def mock_qa(input_dict):
    return {"answer": "This is a placeholder response."}
# # Display conversation history using Streamlit messages
# def display_conversation(history):
   
#     with st.container(border= True): #, height=480):
#         for i in range(len(history["generated2"])):
#             message(
#                 history["past2"][i],
#                 is_user=True,
#                 key=str(i) + "_user",
#             )
#             message(history["generated2"][i], key=str(i), avatar_style="initials", seed="AI")

def save_uploadedfile(uploadedfile):
    
    # with open(os.path.join(r"C:\Users\CVDIV\Downloads\ipf-assistant-devcopy\IPFChat-main\chatBot\user_pdf",uploaded_file.name),"wb") as f:
    with open(os.path.join(r"C:\Users\venkagv\Documents\chatwithdoc-main\chatwithdoc-main\chatBot\DBS",uploadedfile.name),"wb") as f:
        f.write(uploadedfile.getbuffer())
    # return st.success(r"Saved File:{} to C:\Users\CVDIV\Downloads\ipf-assistant-devcopy\IPFChat-main\chatBot\user_pdf".format(uploadedfile.name))
    return st.success("Saved file!")

@st.cache_resource
def declare():
    if "globalFiles" not in st.session_state:
        st.session_state.globalFiles= {}
    return st.session_state.globalFiles

# this is for sessioned(private files)
if "files_dic" not in st.session_state:
    st.session_state.files_dic= {}

# this is for docVisiblity
# values: local, global
if "docVisibility" not in st.session_state:
    st.session_state.docVisibility= ""

# original input container function here->
def chatInterface(qa):

    
    input_container = st.container()
    with input_container:
        
        user_input = st.text_input("", placeholder="Enter your prompt here", key="input")
        

    input_container.float("bottom: 1px;")

    # Initialize session state for generated2 responses and past2 messages
    if "generated2" not in st.session_state:
        st.session_state["generated2"] = ["Ask away any questions you have about the Document."]
    if "past2" not in st.session_state:
        st.session_state["past2"] = ["Hey there!"]

    # to maintain the chat history
    chat_history = []

    # Search the database for a response based on user input and update session state
    if user_input:
        # output = qa({"question": user_input, "chat_history": chat_history})
        output = qa({"query": user_input, "chat_history": chat_history})
        # update chat history by adding this query and answer
        chat_history.append((user_input, output['result']))
        # chat_history.append((user_input, output['answer']))
        print(output["source_documents"])
        #adding the current query and response to appropriate lists
        st.session_state.past2.append(user_input)
        response = str(output["result"])
        # response = str(output["answer"])
        st.session_state.generated2.append(response)

    # Display conversation history using Streamlit messages
    if st.session_state["generated2"]:
        display_conversation(st.session_state)
        # display_conversation(chat_history)

def downloadFiles(filesList):
    # print(filesList, type(filesList))
    for fileName in filesList:
        # print('path', os.path.join(r"X:\CoPilot\chatWithDocUserFiles", fileName))
        with open(os.path.join(r"X:\Technology\CoPilot\chatWithDocUserFiles", fileName),"rb") as file:
        #->with open(os.path.join(r"C:\Users\chauscr\Documents\chatswithdocs-main\chatBot\userdb", fileName),"rb") as file:
                btn = st.download_button(
                        label= fileName,
                        data= file,
                        file_name= fileName,
                        # mime="image/png"
                    )
    st.write('Download Complete!')

def main():
    
    st.markdown("---")
    st.session_state.globalFiles= declare()
    globalFiles_dic= st.session_state.globalFiles

    print('private files:',st.session_state.files_dic)
    print('public/global files',globalFiles_dic)

    # persist_dir= ''
    # db= None

    optionsSelectBox= []
    sessionOptions= st.session_state.files_dic.keys()
    # print('$$$' ,len(sessionOptions), type(sessionOptions))
    # globalOPtions= ['tmo1', 'tmp2', 'tmp3']
    globalOPtions= globalFiles_dic.keys()
    if len(sessionOptions)>0: optionsSelectBox.extend(sessionOptions)
    optionsSelectBox.extend(globalOPtions)
    print(optionsSelectBox)

    # display files for download in the sidebar
    with st.sidebar:
        downloadOption = st.selectbox(
            "Select The Files Available Globally For Download",
            # ["tmp1", "tmp2","tmp3","tmp4","tmp5","tmp6","tmp7","tmp8"],
            optionsSelectBox,
            index= None)

        print('download optons: ', downloadOption)
        # st.button("Download", on_click= downloadFiles, args= [downloadOptions])
        if downloadOption:
            with open(os.path.join(r"C:\Users\venkagv\Documents\chatwithdoc-main\chatwithdoc-main\chatBot\DBS", downloadOption),"rb") as file:
                    btn = st.download_button(
                            label= downloadOption,
                            data= file,
                            file_name= downloadOption,
                            # mime="image/png"
                        )




    # #chat with multiple pdfs
    options= st.multiselect(
        "Pick documents to chat with!",
        # ["tmp1", "tmp2","tmp3","tmp4","tmp5","tmp6","tmp7","tmp8"],
        optionsSelectBox,
        default= None)

    print('Option chosen is:', options)


    uploaded_file = st.file_uploader("Choose a file", type=['txt','docx','pdf','html', 'xml', 'bpmn'], key="uploader")

    view_doc = st.sidebar.checkbox("View Document")

    # if view_doc:
    #     if uploaded_file is not None:
    #         file_name = uploaded_file.name
    #         file_extension = os.path.splitext(file_name)[1].lower()
            
    #         with open(file_name, "wb") as f:
    #             f.write(uploaded_file.getbuffer())
            
    #         st.sidebar.success(f"File uploaded: {file_name}")

    if view_doc and uploaded_file is not None:
        f_name = uploaded_file.name
        file_extension = os.path.splitext(f_name)[1].lower()
        
        with open(f_name, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        if file_extension == ".pdf":
            show_pdf(f_name)
        elif file_extension == ".docx":
            show_docx(f_name)
        elif file_extension == ".txt":
            show_txt(f_name)
        elif file_extension == ".xml":
            show_xml(f_name)
        elif file_extension == ".html":
            show_html(f_name)
        elif file_extension == ".bpmn":
            show_bpmn(f_name)
        st.sidebar.success(f"File uploaded: {f_name}")

    # state variable to capture file id, later used to delete the corresponding embeddings
    if "holdFileId" not in st.session_state:
        st.session_state.holdFileId= ""

    # state variable to decide to generate embeddigns or not
    if "generateEmbs" not in st.session_state:
        st.session_state.generateEmbs= 'True'

    if uploaded_file is not None:
        visOption = st.radio(
            "How would you like the visibility of the doc?",
            ["Visible only to me", "Visible to others"],
            index=None,
        )
      
        
        print('you selected: ', visOption)
        if visOption is None:
            # while option is None:
            st.error('Please select the visibility of the doc above to proceed', icon="🚨")
            # while visOption is None:
            #     # pass
            #     # print('.')
            #     if visOption is not None: break

        try:
            if visOption:
                if visOption== "Visible only to me": ## add to
                    st.session_state.files_dic[st.session_state.uploader.name]= st.session_state.uploader.file_id
                else:
                    st.session_state.globalFiles[st.session_state.uploader.name]= st.session_state.uploader.file_id
        except valueError:
            st.error('Please select the visibility of the doc above to proceed', icon="🚨")

        if visOption:
            # store the db in a unique dir
            # persist_dir='user_data_embs/'+ st.session_state.uploader.file_id
            # base_dir= "X:\CoPilot\Embeddings\chatWithDoc"
            base_dir=r"C:\Users\venkagv\Documents\chatwithdoc-main\chatBot\DBS"
            # tmp_str2= "\"+ st.session_state.uploader.file_id
            # persist_dir = base_dir+ tmp_str2
            persist_dir= os.path.join(base_dir, st.session_state.uploader.file_id)

            # if (st.session_state.generateEmbs=='True'):
            print('Value of generateEmbs:', st.session_state.generateEmbs)
            # if generateEmbs[0]:
            if st.session_state.generateEmbs=='True':
                texts = doc_preprocessing(uploaded_file)
                generateEmbeddings(texts, persist_dir)
                # generateEmbs[0]= 'False'
                st.session_state.generateEmbs= 'False'
                print('Value of generateEmbs after:', st.session_state.generateEmbs)
                # upload the file if a global file
                if visOption!= "Visible only to me":
                    save_uploadedfile(uploaded_file)
    
    # @st.cache_resource
    def add_collection(_db_name):   #the var db_name is unhashable- so telling streamlit to not hash it by adding a '_' before the var name
        db_name_data=_db_name._collection.get(include=['documents','metadatas','embeddings'])
        db._collection.add(
            embeddings=db_name_data['embeddings'],
            metadatas=db_name_data['metadatas'],
            documents=db_name_data['documents'],
            ids=db_name_data['ids']
        )
        print('completed adding collection')

    # #displaying the chat interface
    if (options or (uploaded_file and visOption)):
        if options:
            base_dir= "X:\Technology\CoPilot\Embeddings\chatWithDoc"
            #->base_dir=r"C:\Users\chauscr\Documents\chatwithdoc-main\chatBot\DBS"
            first_option= options[0]
            if first_option in st.session_state.files_dic: persist_dir= os.path.join(base_dir, st.session_state.files_dic[first_option])
            else: persist_dir= os.path.join(base_dir, globalFiles_dic[first_option])
            print('path for first option persist dir:', persist_dir)
            db = Chroma(persist_directory=persist_dir, embedding_function=embeddings)
            # tmp= 1
            for option in options[1:]:
                if option in st.session_state.files_dic: persist_dir= os.path.join(base_dir, st.session_state.files_dic[option])
                else: persist_dir= os.path.join(base_dir, globalFiles_dic[option])
                print('path for persist dir:', persist_dir)
                tmp_db= Chroma(persist_directory=persist_dir, embedding_function=embeddings)
                # add_collection(tmp_db)
                tmp_db_data= tmp_db._collection.get(include=['documents','metadatas','embeddings'])
                db._collection.add(
                    embeddings=tmp_db_data['embeddings'],
                    metadatas=tmp_db_data['metadatas'],
                    documents=tmp_db_data['documents'],
                    ids=tmp_db_data['ids']
                )
                print('Added collections for', option)

        else:
            db = Chroma(persist_directory=persist_dir, embedding_function=embeddings)
        qa = search_pdf(db, secret1)

        if options:
            # option_names= ""
            # for option in options:
            #     option_names+= option+", "
            option_names= ', '.join(options)
            st.write("You are chatting with "+ option_names)
        else:
            st.write("You are chatting with "+ st.session_state.uploader.name)
        chatInterface(qa)

       
       #chatInterface(mock_qa)
    if st.button("Reset Chat"):
            st.session_state.past2 = ["Hey there!"]
            st.session_state.generated2 = ["Ask away any questions you have about the Document."]
            st.session_state.chat_history = []
   

if __name__ == "__main__":
    main()
